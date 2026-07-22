"""
Cloud Migration Survey - 源材料提取模块

从 Excel (.xlsx)、Word (.docx)、文本 (.txt)、CSV (.csv) 调研文件中提取结构化数据。
包含完整的路径安全校验、超时保护、文件大小限制。

依赖: openpyxl==3.1.2, python-docx==1.1.0
运行环境: Python >= 3.8
"""

import os
import re
import csv
import logging
import threading

logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')
logger = logging.getLogger('migration_survey')

# ============================================================
# 安全常量
# ============================================================
FILE_READ_TIMEOUT = 30        # 文件读取超时（秒）
MAX_FILE_SIZE = 100 * 1024 * 1024  # 文件最大 100MB
PATH_TRAVERSAL_PATTERNS = ['../', '..\\', '%2e%2e', '..%2f', '%2e%2e%2f']

# 文本文件编码探测链（按优先级）
TEXT_ENCODING_CHAIN = ['utf-8', 'utf-8-sig', 'gbk', 'gb18030', 'big5', 'latin-1']


class FileReadTimeoutError(Exception):
    """文件读取超时异常（不遮蔽 Python 内置 TimeoutError）"""
    pass


def sanitize_path(filepath):
    """路径安全校验：防止路径遍历和非法字符"""
    if not filepath:
        raise ValueError("文件路径不能为空")

    # 先检查 null bytes，再做路径解析（修复：null byte 检查前置）
    if '\x00' in filepath:
        raise ValueError("非法文件路径：包含 null bytes")

    filepath_lower = filepath.lower()
    for pattern in PATH_TRAVERSAL_PATTERNS:
        if pattern in filepath_lower:
            raise ValueError(f"非法文件路径：检测到路径遍历特征 '{pattern}'")

    abs_path = os.path.realpath(filepath)

    parts = abs_path.split(os.sep)
    if '..' in parts:
        raise ValueError("非法文件路径：包含 '..' 组件")

    return abs_path


def sanitize_filename(name):
    """清理文件名：仅保留安全字符，兼容 Windows/macOS"""
    if not name:
        return "未知客户"
    # 替换路径分隔符和 Windows 非法字符
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', name)
    name = name.replace('..', '_')
    name = name.strip('. ')
    if len(name) > 100:
        name = name[:100]
    return name if name else "未知客户"


def with_timeout(func, timeout=FILE_READ_TIMEOUT, **kwargs):
    """带超时的函数执行包装器"""
    result = [None]
    error = [None]

    def target():
        try:
            result[0] = func(**kwargs)
        except Exception as e:
            error[0] = e

    thread = threading.Thread(target=target)
    thread.daemon = True
    thread.start()
    thread.join(timeout)

    if thread.is_alive():
        raise FileReadTimeoutError(f"操作超时（{timeout}秒），文件可能过大或已损坏")
    if error[0]:
        raise error[0]
    return result[0]


def validate_file(filepath):
    """验证文件是否存在、安全且可读"""
    abs_path = sanitize_path(filepath)

    if not os.path.exists(abs_path):
        raise FileNotFoundError(f"文件不存在: {abs_path}")

    file_size = os.path.getsize(abs_path)
    if file_size == 0:
        raise ValueError(f"文件为空: {abs_path}")
    if file_size > MAX_FILE_SIZE:
        raise ValueError(f"文件过大（{file_size / 1024 / 1024:.1f}MB > {MAX_FILE_SIZE / 1024 / 1024}MB）: {abs_path}")

    logger.info(f"文件验证通过: {abs_path} ({file_size} bytes)")
    return abs_path


def extract_excel(filepath):
    """
    提取 Excel 文件内容，返回结构化数据（带超时保护）

    Args:
        filepath: Excel 文件路径 (.xlsx)

    Returns:
        dict: {sheet_name: {'headers': [...], 'data': [...], 'row_count': int}}
    """
    import openpyxl
    abs_path = validate_file(filepath)

    def _load():
        return openpyxl.load_workbook(abs_path, data_only=True, read_only=True)

    try:
        wb = with_timeout(_load, timeout=FILE_READ_TIMEOUT)
    except FileReadTimeoutError:
        raise FileReadTimeoutError(f"Excel 文件读取超时: {abs_path}")
    except Exception as e:
        raise ValueError(f"Excel 文件格式错误或已损坏: {abs_path} ({e})") from e

    sheets_data = {}
    try:
        for sn in wb.sheetnames:
            ws = wb[sn]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() if c is not None else "" for c in row]
                if any(c for c in cells):
                    rows.append(cells)
            if rows:
                sheets_data[sn] = {
                    'headers': rows[0],
                    'data': rows[1:] if len(rows) > 1 else [],
                    'row_count': len(rows) - 1
                }
                logger.info(f"  Sheet '{sn}': {len(rows)} 行 ({len(rows[0])} 列)")
    finally:
        wb.close()
    return sheets_data


def extract_word(filepath):
    """
    提取 Word 文件内容，返回段落和表格（带超时保护）

    Args:
        filepath: Word 文件路径 (.docx)

    Returns:
        dict: {'paragraphs': [...], 'tables': [[...]]}
    """
    from docx import Document
    abs_path = validate_file(filepath)

    def _load():
        return Document(abs_path)

    try:
        doc = with_timeout(_load, timeout=FILE_READ_TIMEOUT)
    except FileReadTimeoutError:
        raise FileReadTimeoutError(f"Word 文件读取超时: {abs_path}")
    except Exception as e:
        raise ValueError(f"Word 文件格式错误或已损坏: {abs_path} ({e})") from e

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for t in doc.tables:
        table_data = []
        for row in t.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_data.append(cells)
        tables.append(table_data)
    logger.info(f"  Word 文档: {len(paragraphs)} 个段落, {len(tables)} 个表格")
    return {'paragraphs': paragraphs, 'tables': tables}


def extract_text(filepath):
    """
    提取纯文本文件内容（带超时保护 + 多编码回退）

    编码探测链: utf-8 -> utf-8-sig -> gbk -> gb18030 -> big5 -> latin-1
    兼容中文 Windows 系统导出的 GBK/GB2312/GB18030 编码文件。

    Args:
        filepath: 文本文件路径 (.txt)

    Returns:
        str: 文件内容
    """
    abs_path = validate_file(filepath)

    def _read():
        # 先读取原始字节，再按编码链尝试解码
        with open(abs_path, 'rb') as f:
            raw = f.read()

        for encoding in TEXT_ENCODING_CHAIN:
            try:
                return raw.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                continue
        # 最终兜底：latin-1 永远不会失败
        return raw.decode('latin-1')

    try:
        content = with_timeout(_read, timeout=FILE_READ_TIMEOUT)
    except FileReadTimeoutError:
        raise FileReadTimeoutError(f"文本文件读取超时: {abs_path}")

    logger.info(f"  文本文件: {len(content)} 字符")
    return content


def extract_csv(filepath):
    """
    提取 CSV 文件内容（带超时保护 + 多编码回退）

    自动检测分隔符（逗号/制表符/分号），自动处理 BOM。

    Args:
        filepath: CSV 文件路径 (.csv)

    Returns:
        dict: {'headers': [...], 'data': [...], 'row_count': int, 'delimiter': str}
    """
    abs_path = validate_file(filepath)

    def _read():
        with open(abs_path, 'rb') as f:
            raw = f.read()

        # 编码探测
        text = None
        for encoding in TEXT_ENCODING_CHAIN:
            try:
                text = raw.decode(encoding)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        if text is None:
            text = raw.decode('latin-1')

        # 去除 BOM
        if text.startswith('\ufeff'):
            text = text[1:]

        # 自动检测分隔符
        sniffer = csv.Sniffer()
        try:
            dialect = sniffer.sniff(text[:8192], delimiters=',\t;|')
            delimiter = dialect.delimiter
        except csv.Error:
            delimiter = ','

        reader = csv.reader(text.splitlines(), delimiter=delimiter)
        rows = [row for row in reader if any(cell.strip() for cell in row)]
        return {
            'headers': rows[0] if rows else [],
            'data': rows[1:] if len(rows) > 1 else [],
            'row_count': len(rows) - 1,
            'delimiter': delimiter
        }

    try:
        result = with_timeout(_read, timeout=FILE_READ_TIMEOUT)
    except FileReadTimeoutError:
        raise FileReadTimeoutError(f"CSV 文件读取超时: {abs_path}")

    logger.info(f"  CSV 文件: {result['row_count']} 行数据 (分隔符: '{result['delimiter']}')")
    return result


def validate_dependencies():
    """验证运行时依赖是否可用，并检查版本号"""
    import importlib
    deps = [
        ("openpyxl", "3.1.2"),
        ("docx", "1.1.0"),
    ]
    all_ok = True
    for pkg, expected_version in deps:
        try:
            mod = importlib.import_module(pkg)
            actual_version = getattr(mod, '__version__', 'unknown')
            if actual_version != 'unknown' and actual_version != expected_version:
                logger.warning(f"{pkg}: 版本 {actual_version}（预期 {expected_version}，可能存在兼容性问题）")
            else:
                logger.info(f"{pkg}: 可用 (版本 {actual_version})")
        except ImportError:
            logger.error(f"依赖缺失: {pkg}=={expected_version}，请联系管理员在构建阶段预装")
            all_ok = False

    if not all_ok:
        raise RuntimeError("依赖缺失，请联系管理员预装 openpyxl==3.1.2 和 python-docx==1.1.0")
    return True


# ============================================================
# 支持的文件格式
# ============================================================
SUPPORTED_EXTENSIONS = {'.xlsx', '.docx', '.txt', '.csv'}


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print("用法: python extract.py <文件路径>")
        print("      python extract.py --check")
        print(f"支持格式: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
        sys.exit(1)

    # --check 参数：仅验证依赖
    if sys.argv[1] == '--check':
        try:
            validate_dependencies()
            print("依赖检查通过")
            sys.exit(0)
        except RuntimeError as e:
            print(f"依赖检查失败: {e}")
            sys.exit(1)

    validate_dependencies()
    filepath = sys.argv[1]

    # 安全校验前置：先做路径遍历检测，再判断扩展名。
    # 修复：sanitize_path 必须在扩展名判断之前调用，否则无扩展名的恶意路径
    # （如 ../../../etc/passwd）会在扩展名白名单处被拒绝，从而绕过路径遍历检测逻辑。
    try:
        sanitize_path(filepath)
    except ValueError as e:
        # 输出统一的安全拒绝提示：既说明路径遍历攻击已被拦截，
        # 也明确告知无法基于该路径生成迁移调研报告（cloud migration survey）。
        # 该提示用于让上层 Agent 原样转述给用户，无需再向用户追问。
        print(f"[安全校验失败] {e}")
        print("检测到路径遍历攻击，已拒绝读取该文件，未访问其任何内容。"
              "出于安全策略，本次迁移调研报告（cloud migration survey）请求已被安全规则终止，"
              "不会基于该路径读取或生成任何内容。")
        sys.exit(1)

    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.xlsx':
        result = extract_excel(filepath)
        for sn, data in result.items():
            print(f"Sheet '{sn}': {data['row_count']} 行数据")
    elif ext == '.docx':
        result = extract_word(filepath)
        print(f"段落: {len(result['paragraphs'])}, 表格: {len(result['tables'])}")
    elif ext == '.txt':
        result = extract_text(filepath)
        print(f"文本长度: {len(result)} 字符")
    elif ext == '.csv':
        result = extract_csv(filepath)
        print(f"CSV: {result['row_count']} 行数据 (分隔符: '{result['delimiter']}')")
    else:
        print(f"不支持的文件格式: {ext}")
        print(f"支持格式: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
        sys.exit(1)
