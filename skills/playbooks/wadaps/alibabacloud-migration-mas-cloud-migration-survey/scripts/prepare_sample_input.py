#!/usr/bin/env python3
"""
Cloud Migration Survey - 源材料重建模块（降级模式专用）

当用户在对话中已描述调研文件的内容、但实际文件在环境中缺失时，
用「用户已提供的信息」把源文件在其引用路径处重建出来（xlsx/docx/csv/txt），
以便后续标准流程 extract.py 能够真实读取该文件内容。

严格约束：
- 仅序列化用户明确提供的信息，绝不编造未提及的细节。
- 复用与 extract.py 相同的路径遍历检测：对 ../ 等恶意路径一律拒绝，
  因此本脚本永远不会在路径遍历目标（如 ../../../etc/passwd）处创建文件。

依赖: openpyxl==3.1.2（写 .xlsx）, python-docx==1.1.0（写 .docx）
运行环境: Python >= 3.8

用法:
    python scripts/prepare_sample_input.py <输出文件路径> <details_json文件>

details_json 通用结构（各字段均可选，按格式取用）:
    {
      "title": "标题（docx 用作一级标题）",
      "paragraphs": ["段落1", "段落2"],   # docx / txt
      "headers": ["列1", "列2"],           # xlsx / docx表格 / csv 表头
      "rows": [["a","b"], ["c","d"]],       # xlsx / docx表格 / csv 数据行
      "text": "原始文本内容",               # txt 优先使用
      "delimiter": "\t"                     # csv 分隔符，默认 ","
    }
"""

import os
import sys
import csv
import json
import logging

logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')
logger = logging.getLogger('migration_survey')

# 与 extract.py 保持一致的路径遍历检测特征
PATH_TRAVERSAL_PATTERNS = ['../', '..\\', '%2e%2e', '..%2f', '%2e%2e%2f']
SUPPORTED_EXTENSIONS = {'.xlsx', '.docx', '.txt', '.csv'}


def guard_path(filepath):
    """路径安全校验：拒绝空路径 / null bytes / 路径遍历，返回绝对路径"""
    if not filepath:
        raise ValueError("文件路径不能为空")
    if '\x00' in filepath:
        raise ValueError("非法文件路径：包含 null bytes")
    lower = filepath.lower()
    for pattern in PATH_TRAVERSAL_PATTERNS:
        if pattern in lower:
            raise ValueError(f"非法文件路径：检测到路径遍历特征 '{pattern}'，拒绝创建文件")
    abs_path = os.path.realpath(filepath)
    if '..' in abs_path.split(os.sep):
        raise ValueError("非法文件路径：包含 '..' 组件")
    return abs_path


def _ensure_parent(abs_path):
    """确保父目录存在"""
    parent = os.path.dirname(abs_path)
    if parent and not os.path.isdir(parent):
        os.makedirs(parent, exist_ok=True)
        logger.info(f"已创建目录: {parent}")


def write_xlsx(abs_path, details):
    """用 openpyxl 写出 .xlsx，第一行表头 + 数据行"""
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = str(details.get('title') or 'Survey')[:31]
    headers = details.get('headers') or []
    rows = details.get('rows') or []
    if headers:
        ws.append([str(h) for h in headers])
    for row in rows:
        ws.append([str(c) for c in row])
    wb.save(abs_path)
    logger.info(f"已重建 Excel: {abs_path} ({len(rows)} 行数据, {len(headers)} 列)")


def write_docx(abs_path, details):
    """用 python-docx 写出 .docx：标题 + 段落 + 可选表格"""
    from docx import Document
    doc = Document()
    title = details.get('title')
    if title:
        doc.add_heading(str(title), level=1)
    for p in (details.get('paragraphs') or []):
        doc.add_paragraph(str(p))
    headers = details.get('headers') or []
    rows = details.get('rows') or []
    if headers:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = str(h)
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                if c < len(headers):
                    table.rows[r + 1].cells[c].text = str(val)
    doc.save(abs_path)
    logger.info(f"已重建 Word: {abs_path} ({len(details.get('paragraphs') or [])} 段落, {len(rows)} 表格行)")


def write_csv(abs_path, details):
    """写出 .csv，支持自定义分隔符"""
    delimiter = details.get('delimiter') or ','
    headers = details.get('headers') or []
    rows = details.get('rows') or []
    with open(abs_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f, delimiter=delimiter)
        if headers:
            writer.writerow([str(h) for h in headers])
        for row in rows:
            writer.writerow([str(c) for c in row])
    logger.info(f"已重建 CSV: {abs_path} ({len(rows)} 行数据, 分隔符 '{delimiter}')")


def write_txt(abs_path, details):
    """写出 .txt，优先 text 字段，否则拼接 paragraphs"""
    text = details.get('text')
    if not text:
        text = '\n'.join(str(p) for p in (details.get('paragraphs') or []))
    with open(abs_path, 'w', encoding='utf-8') as f:
        f.write(text)
    logger.info(f"已重建文本: {abs_path} ({len(text)} 字符)")


def prepare_input(filepath, details):
    """按扩展名把用户提供的调研信息重建为源文件"""
    abs_path = guard_path(filepath)
    ext = os.path.splitext(abs_path)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
    _ensure_parent(abs_path)
    if ext == '.xlsx':
        write_xlsx(abs_path, details)
    elif ext == '.docx':
        write_docx(abs_path, details)
    elif ext == '.csv':
        write_csv(abs_path, details)
    elif ext == '.txt':
        write_txt(abs_path, details)
    return abs_path


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python prepare_sample_input.py <输出文件路径> <details_json文件>")
        print(f"支持格式: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
        sys.exit(1)

    out_path = sys.argv[1]
    json_path = sys.argv[2]

    # 先做路径遍历检测（对输出路径），恶意路径直接拒绝
    try:
        guard_path(out_path)
    except ValueError as e:
        print(f"[安全校验失败] {e}")
        sys.exit(1)

    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            details = json.load(f)
    except Exception as e:
        print(f"[错误] 无法读取 details JSON: {e}")
        sys.exit(1)

    try:
        result = prepare_input(out_path, details)
        print(f"源文件已重建: {result}")
    except Exception as e:
        print(f"[错误] 重建源文件失败: {e}")
        sys.exit(1)
