"""
Cloud Migration Survey - 报告生成模块

从结构化调研数据生成迁云调研分析报告 (.docx)。
包含完整的样式配置、表格生成、章节排版、质量检查功能。

依赖: python-docx==1.1.0
运行环境: Python >= 3.8
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from datetime import datetime
import os
import re
import json
import logging

logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')
logger = logging.getLogger('migration_survey')


# ============================================================
# 安全常量与工具函数
# ============================================================

# 路径遍历检测模式
PATH_TRAVERSAL_PATTERNS = ['../', '..\\', '%2e%2e', '..%2f', '%2e%2e%2f']
MAX_OUTPUT_FILE_SIZE = 50 * 1024 * 1024  # 输出文件最大 50MB
MAX_INPUT_FILE_SIZE = 50 * 1024 * 1024   # 输入文件最大 50MB
ALLOWED_INPUT_EXTENSIONS = {'.json'}
ALLOWED_OUTPUT_EXTENSIONS = {'.docx'}

# 敏感信息检测正则（用于 quality_check 兜底检测）
SENSITIVE_PATTERNS = {
    '阿里云 AccessKey ID': re.compile(r'LTAI[A-Za-z0-9]{12,20}'),
    '阿里云 AccessKey Secret': re.compile(r'(?<![A-Za-z0-9])[A-Za-z0-9]{30,}(?![A-Za-z0-9/+=])'),  # 30+ 字母数字且不以 base64 填充符结尾
    'AWS Access Key': re.compile(r'AKIA[0-9A-Z]{16}'),
    'AWS Secret Key': re.compile(r'(?i)(aws_secret_access_key|secret_key)\s*[:=]\s*\S+'),
    '通用密码/密钥字段': re.compile(r'(?i)(password|passwd|secret|token|api_key|apikey)\s*[:=]\s*\S+'),
    'IPv4 地址': re.compile(r'\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b'),
    'RSA/SSH 私钥': re.compile(r'-----BEGIN\s+(RSA|DSA|EC|OPENSSH)\s+PRIVATE\s+KEY-----'),
    'JWT Token': re.compile(r'eyJ[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{10,}'),
}

# IPv4 白名单：文档中常见的非敏感 IP（如 0.0.0.0, 127.x, RFC1918 示例等）
IP_EXAMPLE_RANGES = [
    re.compile(r'^127\.'),      # loopback
    re.compile(r'^192\.168\.'), # RFC1918
    re.compile(r'^10\.'),       # RFC1918
    re.compile(r'^172\.(1[6-9]|2\d|3[01])\.'),  # RFC1918
    re.compile(r'^0\.0\.0\.0$'),
    re.compile(r'^255\.255\.255\.'),
]


def _is_example_ip(ip_str):
    """判断 IP 是否为 RFC1918/loopback 等非敏感地址"""
    return any(pat.match(ip_str) for pat in IP_EXAMPLE_RANGES)


def validate_input_path(filepath):
    """
    校验输入文件路径安全性

    检查项：
    - 路径遍历攻击
    - 文件是否存在
    - 文件扩展名白名单
    - 文件大小限制

    Raises:
        ValueError: 路径校验失败
        FileNotFoundError: 文件不存在
    """
    if not filepath:
        raise ValueError("文件路径不能为空")

    # 路径遍历检测
    lower_path = filepath.lower()
    for pattern in PATH_TRAVERSAL_PATTERNS:
        if pattern in lower_path:
            raise ValueError(f"检测到路径遍历攻击模式: {pattern}")

    # 解析为绝对路径
    abs_path = os.path.realpath(filepath)

    # 文件存在性检查
    if not os.path.isfile(abs_path):
        raise FileNotFoundError(f"文件不存在: {abs_path}")

    # 扩展名白名单
    _, ext = os.path.splitext(abs_path)
    if ext.lower() not in ALLOWED_INPUT_EXTENSIONS:
        raise ValueError(f"不允许的文件类型: {ext}，仅支持 {ALLOWED_INPUT_EXTENSIONS}")

    # 文件大小检查
    file_size = os.path.getsize(abs_path)
    if file_size > MAX_INPUT_FILE_SIZE:
        raise ValueError(f"文件过大: {file_size / 1024 / 1024:.1f}MB > {MAX_INPUT_FILE_SIZE / 1024 / 1024:.0f}MB")

    return abs_path


def validate_output_dir(output_dir):
    """
    校验输出目录路径安全性

    检查项：
    - 路径遍历攻击
    - 拒绝写入系统关键目录

    Raises:
        ValueError: 路径校验失败
    """
    if not output_dir:
        raise ValueError("输出目录不能为空")

    # 路径遍历检测
    lower_path = output_dir.lower()
    for pattern in PATH_TRAVERSAL_PATTERNS:
        if pattern in lower_path:
            raise ValueError(f"检测到路径遍历攻击模式: {pattern}")

    abs_path = os.path.realpath(output_dir)

    # 拒绝写入系统关键目录（同时检查原始路径和 realpath，处理 macOS /etc → /private/etc 软链）
    system_dirs = ['/etc', '/usr', '/bin', '/sbin', '/boot', '/sys', '/proc',
                   '/System', '/Library', '/private/etc', '/private/var',
                   'C:\\Windows', 'C:\\Program Files', 'C:\\Program Files (x86)']
    for sys_dir in system_dirs:
        if abs_path.startswith(os.path.realpath(sys_dir)) or output_dir.startswith(sys_dir):
            raise ValueError(f"禁止写入系统目录: {abs_path}")

    return abs_path


def detect_sensitive_data(text, context='报告'):
    """
    检测文本中的敏感信息，返回发现列表

    Args:
        text: 待检测文本
        context: 上下文描述（用于日志）

    Returns:
        list[str]: 发现的敏感信息类型列表
    """
    findings = []
    for name, pattern in SENSITIVE_PATTERNS.items():
        matches = pattern.findall(text)
        if not matches:
            continue

        # 对 IP 地址做白名单过滤
        if name == 'IPv4 地址':
            real_ips = []
            for m in pattern.finditer(text):
                ip = m.group()
                if not _is_example_ip(ip):
                    real_ips.append(ip)
            if real_ips:
                findings.append(f"{name} ({len(real_ips)} 个非内网IP: {', '.join(real_ips[:3])}{'...' if len(real_ips) > 3 else ''})")
        else:
            findings.append(f"{name} ({len(matches)} 处)")

    return findings


# ============================================================
# 配置常量
# ============================================================

class ReportConfig:
    """报告全局配置"""
    PRIMARY_COLOR = '2F5496'
    HEADER_BG = '2F5496'
    HEADER_COLOR = 'FFFFFF'
    ALT_ROW_BG = 'F2F7FF'
    TITLE_COLOR = '2F5496'
    WARNING_BG = 'FFF3E0'
    SUCCESS_BG = 'E8F5E9'

    FONT_CN = '微软雅黑'
    FONT_EN = 'Calibri'

    SIZE_TITLE = 36
    SIZE_SUBTITLE = 28
    SIZE_H1 = 16
    SIZE_H2 = 14
    SIZE_H3 = 12
    SIZE_BODY = 10.5
    SIZE_TABLE = 9
    SIZE_CAPTION = 8

    MARGIN = Cm(2.5)

    MIN_CHAPTERS = 5
    MIN_TABLES = 4
    MAX_EMPTY_CELLS_RATIO = 0.5
    MAX_PLACEHOLDER_RATIO = 0.6  # "待确认" 占比上限


# ============================================================
# 工具函数
# ============================================================

def hex_to_rgb(color):
    """将 HEX 颜色字符串转为 RGBColor 对象"""
    return RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16))


def sanitize_filename(name):
    """清理文件名：仅保留安全字符，兼容 Windows/macOS"""
    if not name:
        return "未知客户"
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', name)
    name = name.replace('..', '_')
    name = name.strip('. ')
    if len(name) > 100:
        name = name[:100]
    return name if name else "未知客户"


def set_cell_shading(cell, color):
    """设置单元格背景色（替换已有 shading）"""
    tcPr = cell._tc.get_or_add_tcPr()
    # 先移除已有的 w:shd 元素，避免重复
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tcPr.append(shading_elm)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, bold=False, size=9, color='000000', align='left', bg=None, font_cn=None, font_en=None):
    """设置单元格文本，自动清空原有 runs"""
    p = cell.paragraphs[0]
    for run in p.runs:
        run.text = ""
        run.font.size = None
        run.font.color.rgb = None
        run.font.name = None
        run.bold = False
        run.italic = False
    if p.runs:
        rPr = p.runs[0]._element.rPr
        if rPr is not None:
            rPr.clear()

    run = p.add_run(str(text) if text else "\u2014")
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = hex_to_rgb(color)
    run.font.name = font_en or ReportConfig.FONT_EN
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn or ReportConfig.FONT_CN)

    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

    if bg:
        set_cell_shading(cell, bg)

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_row_height(row, height_pt):
    """设置行高"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{height_pt}" w:hRule="atLeast"/>')
    trPr.append(trHeight)


def add_styled_table(doc, headers, data, col_widths=None, header_bg=None, header_color=None):
    """创建带样式的表格，返回 table 对象"""
    header_bg = header_bg or ReportConfig.HEADER_BG
    header_color = header_color or ReportConfig.HEADER_COLOR

    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=ReportConfig.SIZE_TABLE,
                      color=header_color, align='center', bg=header_bg)
    set_row_height(table.rows[0], 22)

    for r, row_data in enumerate(data):
        bg = ReportConfig.ALT_ROW_BG if r % 2 == 0 else 'FFFFFF'
        for c, val in enumerate(row_data):
            if c < len(headers):
                set_cell_text(table.rows[r + 1].cells[c], val, size=ReportConfig.SIZE_TABLE, bg=bg)
        set_row_height(table.rows[r + 1], 18)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table


def add_heading_styled(doc, text, level=1):
    """创建带样式的标题"""
    size_map = {1: ReportConfig.SIZE_H1, 2: ReportConfig.SIZE_H2, 3: ReportConfig.SIZE_H3}
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(size_map.get(level, ReportConfig.SIZE_BODY))
        run.font.color.rgb = hex_to_rgb(ReportConfig.TITLE_COLOR)
        run.font.name = ReportConfig.FONT_EN
        run._element.rPr.rFonts.set(qn('w:eastAsia'), ReportConfig.FONT_CN)
    return h


def add_bullet_list(doc, items):
    """添加项目符号列表"""
    last_p = None
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(ReportConfig.SIZE_BODY)
            run.font.name = ReportConfig.FONT_EN
            run._element.rPr.rFonts.set(qn('w:eastAsia'), ReportConfig.FONT_CN)
        last_p = p
    return last_p


def _add_page_field(paragraph):
    """在段落中插入 Word PAGE 域代码"""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar_end)


def _add_toc_field(doc):
    """在文档当前位置插入目录（TOC 域代码）"""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar_begin)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar_separate)

    run4 = p.add_run('（请在 Word 中右键点击此处，选择"更新域"以生成目录）')
    run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run4.font.size = Pt(ReportConfig.SIZE_CAPTION)

    run5 = p.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run5._element.append(fldChar_end)


# ============================================================
# 报告生成主函数
# ============================================================

def generate_report(survey_data, output_dir, customer_name=None, project_name=None):
    """
    生成迁云调研分析报告

    Args:
        survey_data: dict，包含从源材料提取的结构化数据
        output_dir: str，输出目录
        customer_name: str，客户名称（覆盖 survey_data 中的值）
        project_name: str，项目名称

    Returns:
        str: 生成的文件路径
    """
    name = customer_name or survey_data.get('customer_name', '未知客户')
    project = project_name or survey_data.get('project_name', f'{name}云平台迁移项目')
    source_cloud = survey_data.get('source_cloud', '未知')
    today = datetime.now()
    date_str = f'{today.year}年{today.month}月'
    doc_version = survey_data.get('version', 'V1.0')

    logger.info(f"开始生成报告: {name} - {project}")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = ReportConfig.FONT_EN
    style._element.rPr.rFonts.set(qn('w:eastAsia'), ReportConfig.FONT_CN)
    style.font.size = Pt(ReportConfig.SIZE_BODY)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # 设置页边距和页脚（修复：先 unlink 再写页脚）
    for section in doc.sections:
        section.top_margin = ReportConfig.MARGIN
        section.bottom_margin = ReportConfig.MARGIN
        section.left_margin = ReportConfig.MARGIN
        section.right_margin = ReportConfig.MARGIN
        footer = section.footer
        footer.is_linked_to_previous = False  # 先解除链接，再写入页脚
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'\u2014 {name} 云平台迁移调研分析报告 \u2014  第 ')
        run.font.size = Pt(ReportConfig.SIZE_CAPTION)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        _add_page_field(p)  # 插入真正的 PAGE 域代码
        run = p.add_run(' 页')
        run.font.size = Pt(ReportConfig.SIZE_CAPTION)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ========== 封面 ==========
    for _ in range(5):
        doc.add_paragraph()

    h = doc.add_heading(name, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.size = Pt(ReportConfig.SIZE_TITLE)
        run.font.color.rgb = hex_to_rgb(ReportConfig.PRIMARY_COLOR)
        run.font.name = ReportConfig.FONT_EN
        run._element.rPr.rFonts.set(qn('w:eastAsia'), ReportConfig.FONT_CN)

    h2 = doc.add_heading('云平台迁移调研分析报告', level=2)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h2.runs:
        run.font.size = Pt(ReportConfig.SIZE_SUBTITLE)
        run.font.color.rgb = hex_to_rgb(ReportConfig.PRIMARY_COLOR)
        run.font.name = ReportConfig.FONT_EN
        run._element.rPr.rFonts.set(qn('w:eastAsia'), ReportConfig.FONT_CN)

    doc.add_paragraph()

    info_items = [
        f'项目名称：{project}',
        f'客户单位：{name}',
        f'源云平台：{source_cloud}',
        '编制单位：阿里云智能集团',
        f'编制日期：{date_str}',
        f'文档版本：{doc_version}',
        '密    级：内部保密'
    ]
    for line in info_items:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = ReportConfig.FONT_EN
            run._element.rPr.rFonts.set(qn('w:eastAsia'), ReportConfig.FONT_CN)

    doc.add_page_break()

    # ========== 目录 ==========
    add_heading_styled(doc, '目录', level=1)
    _add_toc_field(doc)
    doc.add_page_break()

    # ========== 一、项目概述 ==========
    add_heading_styled(doc, '一、项目概述', level=1)

    add_heading_styled(doc, '1.1 客户概况', level=2)
    overview_data = [
        ['客户名称', name],
        ['源云平台', source_cloud],
        ['目标平台', survey_data.get('target_cloud', '阿里云')],
        ['业务类型', survey_data.get('business_type', '待确认')],
        ['项目阶段', survey_data.get('project_stage', '调研阶段')],
    ]
    add_styled_table(doc, ['项目', '内容'], overview_data, col_widths=[1.5, 4.5])

    add_heading_styled(doc, '1.2 业务特征', level=2)
    biz_items = [
        f"核心场景：{', '.join(survey_data.get('core_scenarios', ['待确认']))}",
        f"高峰期时间：{survey_data.get('peak_hours', '待确认')}",
        f"可用性要求：{survey_data.get('availability_requirement', '待确认')}",
        f"数据规模：{survey_data.get('data_scale', '待确认')}",
        f"迁移窗口：{survey_data.get('migration_window', '待确认')}",
        f"停机容忍度：{survey_data.get('downtime_tolerance', '待确认')}",
    ]
    add_bullet_list(doc, biz_items)

    security = survey_data.get('security_info')
    if security:
        add_heading_styled(doc, '1.3 安全与合规要求', level=2)
        security_data = [[k, str(v)] for k, v in security.items()]
        add_styled_table(doc, ['安全项', '要求'], security_data, col_widths=[2, 4])

    domain = survey_data.get('domain_info')
    if domain:
        add_heading_styled(doc, '1.4 域名与备案信息', level=2)
        domain_data = [[k, str(v)] for k, v in domain.items()]
        add_styled_table(doc, ['项目', '信息'], domain_data, col_widths=[2, 4])

    doc.add_page_break()

    # ========== 二、现状分析 ==========
    add_heading_styled(doc, '二、现状分析', level=1)

    add_heading_styled(doc, '2.1 资源明细', level=2)
    architecture = survey_data.get('architecture', {})
    if architecture:
        arch_data = []
        for comp_type, info in architecture.items():
            if isinstance(info, dict):
                count = info.get('count', '待确认')
                comp_version = info.get('version', '待确认')  # 修复：避免遮蔽 doc_version
                specs = info.get('specs', '待确认')
                migration_notes = info.get('migration_notes', '')
                supplement = info.get('supplement', '')
                arch_data.append([comp_type, count, comp_version, specs, migration_notes, supplement])
            else:
                arch_data.append([comp_type, str(info), '', '', '', ''])
    else:
        arch_data = [
            ['计算', '从源材料提取', '待确认', '规格待确认', '', ''],
            ['存储', '从源材料提取', '待确认', '容量待确认', '', ''],
            ['数据库', '从源材料提取', '待确认', '待确认', '', ''],
            ['网络', '从源材料提取', '待确认', '待确认', '', ''],
            ['中间件', '从源材料提取', '待确认', '待确认', '', ''],
            ['安全', '从源材料提取', '待确认', '待确认', '', ''],
        ]
    add_styled_table(doc, ['组件类型', '资源数量', '版本', '规格', '迁移注意事项', '补充说明'],
                     arch_data, col_widths=[1.0, 1.0, 1.0, 1.2, 1.2, 1.0])

    add_heading_styled(doc, '2.2 关键依赖与瓶颈', level=2)
    deps = survey_data.get('dependencies', ['待分析'])
    add_bullet_list(doc, deps)

    doc.add_page_break()

    # ========== 三、迁移目标与产品映射 ==========
    add_heading_styled(doc, '三、迁移目标与产品映射', level=1)
    mapping = survey_data.get('product_mapping', [])
    if mapping:
        mapping_data = []
        for item in mapping:
            if len(item) >= 3:
                mapping_data.append([str(item[0]), str(item[1]), str(item[2])])
            elif len(item) == 2:
                mapping_data.append([str(item[0]), str(item[1]), '待确认'])
            else:
                mapping_data.append([str(item[0]), '待确认', '待确认'])
    else:
        mapping_data = [
            ['源产品A', '阿里云对应', '迁移方式'],
            ['源产品B', '阿里云对应', '迁移方式'],
        ]
    add_styled_table(doc, ['源云产品', '阿里云产品', '迁移方式'],
                     mapping_data, col_widths=[2, 2, 2])

    doc.add_page_break()

    # ========== 四、技术方案 ==========
    add_heading_styled(doc, '四、技术方案', level=1)
    strategy = survey_data.get('migration_strategy', {})

    if strategy:
        for idx, (component, details) in enumerate(strategy.items(), 1):  # 修复：O(n^2) → O(n)
            add_heading_styled(doc, f'4.{idx} {component}', level=2)
            if isinstance(details, str):
                doc.add_paragraph(details)
            elif isinstance(details, dict):
                detail_items = [[k, str(v)] for k, v in details.items()]
                add_styled_table(doc, ['项目', '方案'], detail_items, col_widths=[2, 4])
    else:
        add_heading_styled(doc, '4.1 计算迁移策略', level=2)
        doc.add_paragraph('虚拟机部署业务可以采用 SMC 镜像迁移工具进行，但更建议采用重新部署的方式，避免把源主机上各类缓存等异常带入新环境中；如果容器化的业务，建议通过 CI/CD 方式在阿里云容器服务 ACK 中重新部署。')
        add_heading_styled(doc, '4.2 存储迁移策略', level=2)
        doc.add_paragraph('对象存储迁移使用阿里云在线迁移服务进行，需要重点关注元数据的同步（安全策略、生命周期配置等）和增量迁移方式（增量延迟受限于源端文件数量，每次增量扫描时间过长可能导致增量时间不可控，可以考虑增量事件订阅、镜像回源等方式处理增量）。')
        add_heading_styled(doc, '4.3 数据库迁移策略', level=2)
        doc.add_paragraph('常规使用 DTS 进行全量+增量数据同步，割接时短暂停写完成最终切换。但是需要注意如果涉及到版本升级时，建议引入流量回放工具进行兼容性和性能评估（流量回放地址：https://help.aliyun.com/zh/cmh/cloud-migration-hub/traffic-replay），另外版本升级割接时需要考虑反向同步任务可能无法配置，需要提前考虑影响。')
        add_heading_styled(doc, '4.4 网络与安全策略', level=2)
        doc.add_paragraph('规划 VPC/VSwitch 网络架构，确定源端互联方式，配置安全组和 ACL，部署 WAF 和高防等安全防护。需要根据迁移数据量、时间要求评估带宽需求（VPN 还是专线），如果用 VPN 方式需尽量避免跨云调用。VPC 划分时需考虑转发路由器 TR 带来的流量转发成本，大量数据同步场景的业务尽量在一个 VPC。迁移过程中对公网流量无法清晰把握的情况下建议采用共享带宽。')
        add_heading_styled(doc, '4.5 中间件策略', level=2)
        doc.add_paragraph('消息中间件迁移需要考虑端口变化、命名规范差异，阿里云消息中间件产品和开源版本不严格对标，但基本都兼容开源版本。')

    doc.add_page_break()

    # ========== 五、版本兼容性风险（新增章节） ==========
    add_heading_styled(doc, '五、版本兼容性风险', level=1)
    version_risks = survey_data.get('version_risks', [])
    if version_risks:
        risk_rows = []
        for vr in version_risks:
            if isinstance(vr, dict):
                risk_rows.append([
                    vr.get('component', '待确认'),
                    vr.get('source_version', '待确认'),
                    vr.get('target_version', '待确认'),
                    vr.get('risk_level', '待评估'),
                    vr.get('notes', ''),
                ])
            elif isinstance(vr, (list, tuple)) and len(vr) >= 4:
                risk_rows.append([str(x) for x in vr[:5]])
        if risk_rows:
            vr_table = add_styled_table(doc,
                ['组件', '源端版本', '目标版本', '风险等级', '说明'],
                risk_rows, col_widths=[1.2, 1.0, 1.0, 0.8, 2.0])
            # 高风险行标橙
            if vr_table:
                for row in vr_table.rows[1:]:
                    if len(row.cells) >= 4:
                        level = row.cells[3].paragraphs[0].text.strip()
                        if level in ('高', '🔴 高', '🔴'):
                            for cell in row.cells:
                                set_cell_shading(cell, ReportConfig.WARNING_BG)
    else:
        doc.add_paragraph('请在调研过程中确认各组件的源端版本号，并参照 version-risks.md 评估版本兼容性风险。'
                          '重点关注：数据库版本（MySQL/PostgreSQL/MongoDB/Redis）、中间件版本（Kafka/RocketMQ）、'
                          '容器编排版本（K8s API 废弃检查）。')

    doc.add_page_break()

    # ========== 六、风险评估与应对 ==========
    add_heading_styled(doc, '六、风险评估与应对', level=1)
    risks = survey_data.get('risks', [])
    if risks:
        risk_data = []
        for risk in risks:
            if isinstance(risk, (list, tuple)):
                risk_data.append([str(r) for r in risk])
            else:
                risk_data.append([str(risk)])
    else:
        risk_data = [
            ['数据迁移失败', '低', '高', '全量备份 + 增量同步，割接前完成数据校验'],
            ['应用不兼容', '中', '高', '提前 POC 测试，准备回退方案'],
            ['割接时间超出窗口', '中', '高', '制定详细割接计划，预留缓冲时间'],
            ['DNS 切换延迟', '低', '中', '提前调低 TTL 值，准备多线路切换方案'],
        ]
    risk_table = add_styled_table(doc, ['风险项', '概率', '影响', '应对措施'],
                                   risk_data, col_widths=[1.8, 0.8, 0.8, 2.6])
    if risk_table:
        for row in risk_table.rows[1:]:
            if len(row.cells) >= 3:
                impact = row.cells[2].paragraphs[0].text.strip()
                if impact == '高':
                    for cell in row.cells:
                        set_cell_shading(cell, ReportConfig.WARNING_BG)

    doc.add_page_break()

    # ========== 七、待补充信息 ==========
    add_heading_styled(doc, '七、待补充信息', level=1)
    pending = survey_data.get('pending_items', [])
    if pending:
        add_bullet_list(doc, pending)
    else:
        doc.add_paragraph('暂无待补充信息', style='List Bullet')

    # ========== 八、建议下一步行动 ==========
    add_heading_styled(doc, '八、建议下一步行动', level=1)
    next_steps = survey_data.get('next_steps', [
        '导出完整资源清单',
        '完成 POC 技术验证',
        '确认待补充信息',
        '确认各组件版本号并评估兼容性风险',
        '制定详细迁移计划与时间表',
        '召开项目启动会',
    ])
    for i, item in enumerate(next_steps, 1):
        p = doc.add_paragraph(f'{i}. {item}')
        for run in p.runs:
            run.font.size = Pt(ReportConfig.SIZE_BODY)
            run.font.name = ReportConfig.FONT_EN
            run._element.rPr.rFonts.set(qn('w:eastAsia'), ReportConfig.FONT_CN)

    # ========== 保存 ==========
    safe_name = sanitize_filename(name)
    output_filename = f'{safe_name}_迁移调研分析报告.docx'

    # 安全校验：输出目录
    output_dir = validate_output_dir(output_dir)
    output_path = os.path.join(output_dir, output_filename)

    # 文件已存在时记录警告（不静默覆盖）
    if os.path.exists(output_path):
        logger.warning(f"输出文件已存在，将覆盖: {output_path}")

    os.makedirs(output_dir, exist_ok=True)

    try:
        doc.save(output_path)
        logger.info(f"报告已保存: {output_path}")
    except PermissionError:
        raise PermissionError(f"无法写入文件（权限不足或文件被占用）: {output_path}")
    except OSError as e:
        raise OSError(f"文件保存失败: {output_path} ({e})")

    return output_path


# ============================================================
# 质检函数
# ============================================================

def quality_check(filepath):
    """
    对生成的报告进行质量检查

    合格标准：
    - 至少 5 个章节
    - 至少 4 个表格
    - 空单元格 <= 50%
    - "待确认" 占位符 <= 60%
    - 包含必要章节（项目概述/现状分析/产品映射/技术方案/风险）
    - 不包含报价内容（段落 + 表格全扫描）
    - 敏感信息检测（IP/AccessKey/密码/密钥/JWT 等正则兜底）

    Returns:
        dict: {'passed': bool, 'chapters': int, 'tables': int, ...}
    """
    doc = Document(filepath)

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    # 修复：章节检测使用"X、"前缀匹配，避免子标题误匹配
    chapters = [h for h in headings if re.match(r'^[一二三四五六七八九十]+、', h)]

    total_tables = len(doc.tables)
    total_empty = 0
    total_cells = 0
    placeholder_count = 0

    # 收集所有文本（段落 + 表格）用于报价检测和占位符统计
    all_text_parts = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_cells += 1
                text = cell.text.strip()
                all_text_parts.append(text)
                if not text or text == '\u2014':
                    total_empty += 1
                if '待确认' in text:
                    placeholder_count += 1

    for p in doc.paragraphs:
        all_text_parts.append(p.text)

    full_text = '\n'.join(all_text_parts)

    required = ['项目概述', '现状分析', '产品映射', '技术方案', '风险']
    missing = []
    for r in required:
        if not any(r in h for h in headings):
            missing.append(r)

    empty_ratio = total_empty / total_cells if total_cells > 0 else 0
    placeholder_ratio = placeholder_count / total_cells if total_cells > 0 else 0

    results = {
        'passed': True,
        'chapters': len(chapters),
        'tables': total_tables,
        'empty_cells': total_empty,
        'empty_ratio': f'{empty_ratio:.1%}',
        'placeholder_count': placeholder_count,
        'placeholder_ratio': f'{placeholder_ratio:.1%}',
        'missing_chapters': missing,
        'warnings': []
    }

    if len(chapters) < ReportConfig.MIN_CHAPTERS:
        results['passed'] = False
        results['warnings'].append(f"章节数不足: {len(chapters)} < {ReportConfig.MIN_CHAPTERS}")
    if total_tables < ReportConfig.MIN_TABLES:
        results['passed'] = False
        results['warnings'].append(f"表格数不足: {total_tables} < {ReportConfig.MIN_TABLES}")
    if empty_ratio > ReportConfig.MAX_EMPTY_CELLS_RATIO:
        results['passed'] = False
        results['warnings'].append(f"空单元格比例过高: {empty_ratio:.1%} > {ReportConfig.MAX_EMPTY_CELLS_RATIO:.1%}")
    if placeholder_ratio > ReportConfig.MAX_PLACEHOLDER_RATIO:
        results['warnings'].append(f"占位符(待确认)比例过高: {placeholder_ratio:.1%}，报告信息可能不完整")
    if missing:
        results['passed'] = False
        results['warnings'].append(f"缺少必要章节: {', '.join(missing)}")

    # 报价检测：扫描段落 + 表格全文（修复：之前只扫段落）
    price_keywords = ['报价', '单价', '人天费用', '月度费用', '年度费用', '元/月', '元/年']
    found_prices = [kw for kw in price_keywords if kw in full_text]
    if found_prices:
        results['passed'] = False
        results['warnings'].append(f"报告包含报价内容（安全红线违规）: {', '.join(found_prices)}")

    # 敏感信息检测：代码级防护兜底（IP、AccessKey、密码、密钥等）
    sensitive_findings = detect_sensitive_data(full_text)
    if sensitive_findings:
        for finding in sensitive_findings:
            results['warnings'].append(f"[敏感信息检测] 报告中发现 {finding}，请确认是否为公开信息，必要时进行脱敏处理")
        results['sensitive_data_count'] = len(sensitive_findings)

    logger.info(f"质检结果: {'通过' if results['passed'] else '未通过'}")
    logger.info(f"  章节数: {results['chapters']}")
    logger.info(f"  表格数: {results['tables']}")
    logger.info(f"  空单元格: {results['empty_cells']} ({results['empty_ratio']})")
    logger.info(f"  占位符: {results['placeholder_count']} ({results['placeholder_ratio']})")
    if results['warnings']:
        for w in results['warnings']:
            logger.warning(f"  - {w}")

    return results


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print("用法: python generate_report.py <survey_data_json_file> [output_dir]")
        print("  survey_data_json_file: 包含调研数据的 JSON 文件路径")
        sys.exit(1)

    json_path = sys.argv[1]

    # 安全校验：输入文件路径
    try:
        json_path = validate_input_path(json_path)
    except (ValueError, FileNotFoundError) as e:
        print(f"[安全校验失败] {e}")
        sys.exit(1)

    # 读取 JSON（带大小限制和异常处理）
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            survey_data = json.load(f)
    except json.JSONDecodeError as e:
        print(f"[JSON 解析失败] 文件格式错误: {e}")
        sys.exit(1)
    except UnicodeDecodeError:
        print(f"[编码错误] JSON 文件应为 UTF-8 编码")
        sys.exit(1)

    # 安全校验：输出目录
    output_dir = sys.argv[2] if len(sys.argv) > 2 else './output/'
    try:
        output_dir = validate_output_dir(output_dir)
    except ValueError as e:
        print(f"[安全校验失败] {e}")
        sys.exit(1)

    output_path = generate_report(survey_data, output_dir)

    results = quality_check(output_path)
    if not results['passed']:
        print("质检未通过，请检查报告内容:")
        for w in results['warnings']:
            print(f"  - {w}")
        sys.exit(1)
    else:
        print(f"质检通过: {output_path}")
        if results.get('sensitive_data_count', 0) > 0:
            print(f"[注意] 检测到 {results['sensitive_data_count']} 类敏感信息，请查看详情")
