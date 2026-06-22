#!/usr/bin/env python3
"""
迁云方案 Markdown → Word 文档转换工具
基于 python-docx 精确控制 Word 排版格式
推荐使用: uv run --with python-docx python3 md_to_docx.py <input.md> [output.docx]
"""

import sys
import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── 样式配置 ───────────────────────────────────────────────

FONT_CN = "微软雅黑"          # 中文字体（可改为"宋体"、"黑体"等）
FONT_EN = "Calibri"           # 英文字体
FONT_SIZE_BODY = Pt(10.5)     # 正文字号（五号）
FONT_SIZE_TABLE = Pt(9)       # 表格字号（小五号）
LINE_SPACING = 1.5            # 行距
TITLE_COLOR = RGBColor(0x1F, 0x3A, 0x5F)   # 标题颜色（深蓝）
TABLE_HEADER_BG = "1F3A5F"    # 表头背景色（深蓝）
TABLE_HEADER_FG = "FFFFFF"    # 表头文字颜色（白色）
TABLE_STRIPE_BG = "F2F6FA"    # 表格斑马纹背景色
TABLE_BORDER_COLOR = "B0C4DE" # 表格边框颜色


# ─── 字体设置工具 ─────────────────────────────────────────────

def set_run_font(run, font_cn=FONT_CN, font_en=FONT_EN, size=FONT_SIZE_BODY,
                 bold=False, color=None):
    """统一设置 run 的中英文字体"""
    run.font.name = font_en
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置中文字体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)


def set_paragraph_format(paragraph, alignment=None, space_before=None,
                         space_after=None, line_spacing=LINE_SPACING,
                         left_indent=None, first_line_indent=None):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if left_indent is not None:
        pf.left_indent = left_indent
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


# ─── 文档创建 ─────────────────────────────────────────────────

def create_document():
    """创建带预设样式的文档"""
    doc = Document()

    # 页面设置（A4）
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 默认段落样式
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.font.size = FONT_SIZE_BODY
    style.paragraph_format.line_spacing = LINE_SPACING
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    return doc


# ─── 标题处理 ─────────────────────────────────────────────────

def add_title(doc, text):
    """添加文档标题（H1）"""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(12), space_after=Pt(18))
    run = p.add_run(text)
    set_run_font(run, size=Pt(22), bold=True, color=TITLE_COLOR)


def add_heading(doc, text, level):
    """添加各级标题"""
    size_map = {2: Pt(16), 3: Pt(14), 4: Pt(12)}
    space_map = {2: (Pt(18), Pt(10)), 3: (Pt(14), Pt(8)), 4: (Pt(10), Pt(6))}

    p = doc.add_paragraph()
    before, after = space_map.get(level, (Pt(10), Pt(6)))
    set_paragraph_format(p, space_before=before, space_after=after)
    run = p.add_run(text)
    set_run_font(run, size=size_map.get(level, Pt(12)), bold=True,
                 color=TITLE_COLOR)


# ─── 表格处理 ─────────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            element = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="single" '
                f'w:sz="4" w:space="0" w:color="{kwargs[edge]}"/>'
            )
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_bg(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, rows):
    """添加格式化表格"""
    if not rows or len(rows) < 1:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                continue
            cell = table.cell(i, j)
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 写入文本
            p = cell.paragraphs[0]
            set_paragraph_format(p, space_before=Pt(2), space_after=Pt(2),
                                 line_spacing=1.2)
            add_rich_text(p, cell_text, size=FONT_SIZE_TABLE,
                          bold=(i == 0),
                          color=RGBColor(0xFF, 0xFF, 0xFF) if i == 0 else None)

            # 边框
            set_cell_border(cell, top=TABLE_BORDER_COLOR,
                            bottom=TABLE_BORDER_COLOR,
                            left=TABLE_BORDER_COLOR,
                            right=TABLE_BORDER_COLOR)

            # 表头样式
            if i == 0:
                set_cell_bg(cell, TABLE_HEADER_BG)
            # 斑马纹（偶数行）
            elif i % 2 == 0:
                set_cell_bg(cell, TABLE_STRIPE_BG)

    doc.add_paragraph()  # 表后间距


# ─── 富文本处理 ────────────────────────────────────────────────

def add_rich_text(paragraph, text, size=FONT_SIZE_BODY, bold=False, color=None):
    """处理 Markdown 行内标记（**粗体**）"""
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        is_bold = (idx % 2 == 1) or bold
        run = paragraph.add_run(part)
        set_run_font(run, size=size, bold=is_bold, color=color)


# ─── Markdown 解析 ─────────────────────────────────────────────

def parse_table_lines(lines, start_idx):
    """解析 Markdown 表格行"""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # 跳过分隔行 |---|---|
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def convert(md_content, output_path):
    """Markdown → Word 文档转换主函数"""
    doc = create_document()
    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行跳过
        if not stripped:
            i += 1
            continue

        # ── 标题 ──
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 1:
                add_title(doc, text)
            else:
                add_heading(doc, text, level)
            i += 1
            continue

        # ── 表格 ──
        if stripped.startswith('|'):
            rows, i = parse_table_lines(lines, i)
            add_table(doc, rows)
            continue

        # ── 无序列表 ──
        list_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(2)
            p = doc.add_paragraph()
            left = Cm(1.0) if indent < 4 else Cm(2.0)
            set_paragraph_format(p, left_indent=left,
                                 space_before=Pt(2), space_after=Pt(2))
            run = p.add_run("• ")
            set_run_font(run, size=FONT_SIZE_BODY)
            add_rich_text(p, text)
            i += 1
            continue

        # ── 有序列表 ──
        ol_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if ol_match:
            indent = len(ol_match.group(1))
            num = ol_match.group(2)
            text = ol_match.group(3)
            p = doc.add_paragraph()
            left = Cm(1.0) if indent < 4 else Cm(2.0)
            set_paragraph_format(p, left_indent=left,
                                 space_before=Pt(2), space_after=Pt(2))
            run = p.add_run(f"{num}. ")
            set_run_font(run, size=FONT_SIZE_BODY, bold=False)
            add_rich_text(p, text)
            i += 1
            continue

        # ── 引用块 > ──
        quote_match = re.match(r'^>\s*(.*)$', stripped)
        if quote_match:
            text = quote_match.group(1)
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=Cm(1.0),
                                 space_before=Pt(4), space_after=Pt(4))
            # 左侧蓝色竖线效果：用缩进+颜色模拟
            run = p.add_run("│ ")
            set_run_font(run, size=FONT_SIZE_BODY,
                         color=RGBColor(0x64, 0x95, 0xED))
            add_rich_text(p, text, color=RGBColor(0x66, 0x66, 0x66))
            i += 1
            continue

        # ── 分割线 --- ──
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            p = doc.add_paragraph()
            set_paragraph_format(p, space_before=Pt(6), space_after=Pt(6))
            run = p.add_run("─" * 60)
            set_run_font(run, size=Pt(8), color=RGBColor(0xCC, 0xCC, 0xCC))
            i += 1
            continue

        # ── 普通段落 ──
        p = doc.add_paragraph()
        set_paragraph_format(p, space_before=Pt(3), space_after=Pt(3))
        add_rich_text(p, stripped)
        i += 1

    doc.save(output_path)
    return output_path


# ─── 入口 ──────────────────────────────────────────────────────

def _validate_path(path: str, expected_ext: str, label: str) -> str:
    """校验路径安全性：规范化、遍历检测、扩展名检查。"""
    # 规范化路径，消除 .. / ./ 等
    real_path = os.path.realpath(path)

    # 白名单目录：只允许在当前工作目录及其子目录下操作
    allowed_dir = os.path.realpath(os.getcwd())
    if not real_path.startswith(allowed_dir + os.sep) and real_path != allowed_dir:
        print(f"错误: {label}路径不在允许的目录范围内 - {path}")
        print(f"   允许目录: {allowed_dir}")
        sys.exit(1)

    # 扩展名校验
    _, ext = os.path.splitext(real_path)
    if ext.lower() != expected_ext.lower():
        print(f"错误: {label}文件扩展名必须为 {expected_ext}，实际为 {ext} - {path}")
        sys.exit(1)

    return real_path


def main():
    if len(sys.argv) < 2:
        print("用法: uv run --with python-docx python3 md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)

    raw_input = sys.argv[1]
    raw_output = sys.argv[2] if len(sys.argv) >= 3 else os.path.splitext(raw_input)[0] + '.docx'

    # 输入校验：必须是 .md，且在安全目录内
    input_path = _validate_path(raw_input, '.md', '输入')
    # 输出校验：必须是 .docx，且在安全目录内
    output_path = _validate_path(raw_output, '.docx', '输出')

    if not os.path.exists(input_path):
        print(f"错误: 文件不存在 - {input_path}")
        sys.exit(1)

    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    result = convert(md_content, output_path)
    print(f"✅ Word 文档已生成: {result}")
    file_size = os.path.getsize(result)
    print(f"   文件大小: {file_size / 1024:.1f} KB")


if __name__ == '__main__':
    main()
